from docx import Document
import  re
import os
import nltk
from nltk.corpus import brown
from selenium import webdriver
# 时间存储字典
def dict():
    # 打开字典并去掉换行符
    with open("../time.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
        # 去除换行符
        result = ([x.strip() for x in lines if x.strip() != ''])
        # 将整理好字典提取成为全局字典
        for x in result:
            Dict.append(x)
    return Dict
# 地点存储词典
def geodict():
    # 打开字典并去掉换行符
    with open("../geosubstance1.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
        # 去除换行符
        result = ([x.strip() for x in lines if x.strip() != ''])
        # 将整理好字典提取成为全局字典
        for x in result:
            GeoDict.append(x)
    return GeoDict
# 读取每段文本
# 进行初步清理并合并成一个文本
def getParagraphsText(doc,Namedocx):
    paragraphsText = []
    paragraphsText_new = []
    for i in range(0,len(doc.paragraphs)):
        print("这是:",doc.paragraphs[i].text)
    for i in range(0, len(doc.paragraphs)):
        paragraphsText.append(doc.paragraphs[i].text)
    text = ([x.strip() for x in paragraphsText if x.strip() != ''])
    # 初步数据清理从abstract开始删除acknowedge以后字段
    a = 0
    for i in range(len(text)):
        if text[i] == "Abstract":
            a = i
    for i in range(a+1,len(text)):
        if text[i] == 'Acknowledgements':
            break
        else:
            #text[i] = re.sub(u"\\(.*?\\)|\\{.*?}|\\[.*?]", '', text[i])
            text[i] = re.sub(r"\d+\.(\d+)?", '', text[i])
            text[i] = re.sub(r"\d+\.(\d+)? m",'',text[i])
            text[i] = re.sub(r"\d+\.(\d+)? km", '', text[i])
            text[i] = re.sub(r"Fig. +\d", '', text[i])
            text[i] = re.sub(r"Figs. \d and \d", '', text[i])
            text[i] = re.sub(r"Figs. +\d", '', text[i])
            text[i] = re.sub(r" (\w)?\.", '', text[i])
            text[i] = re.sub(r" (\w)?\.(\w)?\.", '', text[i])
            paragraphsText_new.append(text[i].replace("ca.","ca").replace("i.e.","").replace("etc.","etc").replace("e.g.,","").replace("e.g.","").replace("al.","al")\
               .replace("cf.","cf").replace("Fig","").replace("B.P.","").replace("S.E.","").replace(" .","").replace("sp.","").replace("c.","").replace("R.","")
                                      .replace("e.","").replace("C.","").replace("T.","").replace("S.E.","").replace("s.l.","").replace("B.V.","").replace("%","").strip())
    f = open('../txt/' + Namedocx.split('.')[0] + ".txt", "w", encoding='utf-8')
    for line in paragraphsText_new:
        f.write(' '+line)
    print(Namedocx + "保存成功")
    print("清洗文本")
    f.close()
# 进行提取筛选
def filter():
    osfile = []
    dir = r"D:\pythonProjectRelationshipToNeo4j\txt"
    for root, dirs, files in os.walk(dir):
        for file in files:
            if file.split('.')[1] == 'txt':
                osfile.append(file)
    for indexFile in range(len(osfile)):
        splitTxt = []
        splitTxtSpace = []
        filterResult = []
        ResultDoc = []
        with open('../txt/' + osfile[indexFile], "r", encoding="utf-8") as f:
            lines = f.readlines()
            resultTxt = ([x.strip() for x in lines if x.strip() != ''])
            for index1 in range(len(resultTxt)):
               splitTxt=resultTxt[index1].split('.')
               for index2 in range(len(splitTxt)):
                 splitTxtSpace.append(splitTxt[index2].strip())
            splitTxtSpace = [i for i in splitTxtSpace if (len(str(i)) != 0)]
            # 匹配抽取句子
            for i in range(len(splitTxtSpace) - 1):
                flag1geo = 0
                line1geo = 0
                flag1tim = 0
                line1tim = 0
                line2geo = 0
                flag2geo = 0
                line2tim = 0
                flag2tim = 0
                for geo in range(len(GeoDict)):
                    if splitTxtSpace[i].find(GeoDict[geo]) != -1:
                        line1geo = 1
                        flag1geo = i
                        #print(str(line1geo)+" flag1geo: " +str(i)+" geo: "+splitTxtSpace[i])
                        for tim in range(len(Dict)):
                            if splitTxtSpace[i+1].find(Dict[tim]) != -1:
                                line2tim = 1
                                flag2tim = i+1
                                print(str(line1geo) + " flag1geo: " + str(flag1geo) + " geo: " + splitTxtSpace[i])
                                print(str(line2tim) + " flag2tim: " + str(flag2tim) + " tim: " + splitTxtSpace[i + 1])
                                #filterResult.append(splitTxtSpace[i]+'. '+splitTxtSpace[i+1]+'.')
                print("+++++++++++++++++++++++++++++++++++++")
                for tim in range(len(Dict)):
                    if splitTxtSpace[i].find(Dict[tim]) != -1:
                        line1tim = 1
                        flag1tim = i
                        #print(str(line1tim)+" flag1tim: " +str(i)+" tim: "+splitTxtSpace[i])
                        for geo in range(len(GeoDict)):
                            if splitTxtSpace[i+1].find(GeoDict[geo]) != -1:
                                line2geo = 1
                                flag2geo = i+1
                                print(str(line1tim) + " flag1tim: " + str(flag1tim) + " tim: " + splitTxtSpace[i])
                                print(str(line2geo) + " flag2geo: " + str(flag2geo) + " geo: " + splitTxtSpace[i+1])
                                #filterResult.append(splitTxtSpace[i] + '. ' + splitTxtSpace[i + 1] + '.')

                if line1tim ==1 and line1geo ==1:
                    print(str(flag1geo)+" flaggeo "+str(flag1tim)+" flagtim")
                    filterResult.append(splitTxtSpace[i]+'.')
                elif (line1tim == 1 and line2geo == 1) or (line1geo == 1 and line2tim == 1):
                    filterResult.append(splitTxtSpace[i] + '. ' + splitTxtSpace[i + 1] + '.')
            for index3 in range(len(filterResult)):
                print(filterResult[index3])
            print("===============================================")
            seen = set()
            for item in filterResult:
                if item not in seen:
                    seen.add(item)
                    ResultDoc.append(item)

            for index4 in range(len(ResultDoc)):
                print(ResultDoc[index4])
            f.close()

            if ResultDoc:
                f1 = open('../ToNeo/' + osfile[indexFile].split('.')[0] + ".txt", "w", encoding='utf-8')
                for line in range(len(ResultDoc)):
                    f1.write(ResultDoc[line] + '\n')

                print(osfile[indexFile].split('.')[0] + ".txt" + "保存成功")
                f1.close()
                print("okA")
# 抽取
brown_train = brown.tagged_sents(categories='news')
regexp_tagger = nltk.RegexpTagger(
    [(r'^-?[0-9]+(.[0-9]+)?$', 'CD'),
     (r'(-|:|;)$', ':'),
     (r'\'*$', 'MD'),
     (r'(The|the|A|a|An|an)$', 'AT'),
     #(r'.*able$', 'JJ'),
     (r'^±', 'JJ'),
     (r'^[A-Z].*$', 'NNP'),
     (r'.*ness$', 'NN'),
     (r'.*ly$', 'RB'),
     (r'.*s$', 'NNS'),
     (r'^~?[0-9]','VB'),
     (r'.*ing$', 'VBG'),
     (r'.*ed$', 'VBD'),
     #(r'et$','FW-CC'),
     #(r'al$','FW-NN'),
     (r'and$','CC'),
     (r'.*', 'NN')
     ])
unigram_tagger = nltk.UnigramTagger(brown_train, backoff=regexp_tagger)
bigram_tagger = nltk.BigramTagger(brown_train, backoff=unigram_tagger)
# This is our semi-CFG; Extend it according to your own needs
#############################################################################
cfg = {}
cfg["NNP+NNP"] = "NNP"
cfg["CD+NNP"] = "NNP"
cfg["JJ+NNP"] = "NNP"
cfg["VBD+NN"] = "NNP"
cfg["VB+NNP"] = "NNP"
cfg["NN+NN"] = "NNI"
cfg["NNI+NN"] = "NNI"
cfg["CD+JJ"] = "NNP"
cfg["CD+NN"] = "NNP"
cfg["CD+IN"] = "NNP"
cfg["JJ+CD"] = "NNP"
cfg["JJ+JJ"] = "JJ"
cfg["JJ+NN"] = "NNI"
#cfg["FW-CC+FW-NN"] = "NNI"
cfg["CD+CC"] = "NNP"
#############################################################################
class NPExtractor(object):
    def __init__(self, sentence):
        self.sentence = sentence
    def dict(self):
        Dict = []
        with open("../time.txt", "r", encoding="utf-8") as f:
            lines = f.readlines()
            result = ([x.strip() for x in lines if x.strip() != ''])
            for x in result:
                Dict.append(x)
        return Dict

    def geosubstance(self):
        geo = []
        with open("../geosubstance1.txt", "r", encoding="utf-8") as f:
            lines = f.readlines()
            result = ([x.strip() for x in lines if x.strip() != ''])
            for x in result:
                geo.append(x)
        return geo
    def tokenize_sentence(self, sentence):
        tokens = nltk.word_tokenize(sentence)
        return tokens
    def normalize_tags(self, tagged):
        n_tagged = []
        for t in tagged:
            if t[1] == "NP-TL" or t[1] == "NP":
                n_tagged.append((t[0], "NNP"))
                continue
            if t[1].endswith("-TL"):
                n_tagged.append((t[0], t[1][:-3]))
                continue
            if t[1].endswith("S"):
                n_tagged.append((t[0], t[1][:-1]))
                continue
            n_tagged.append((t[0], t[1]))
        return n_tagged
    def extract(self):
        tokens = self.tokenize_sentence(self.sentence)
        tags = self.normalize_tags(bigram_tagger.tag(tokens))
        merge = True
        while merge:
            merge = False
            for x in range(0, len(tags) - 1):
                t1 = tags[x]
                t2 = tags[x + 1]
                key = "%s+%s" % (t1[1], t2[1])
                value = cfg.get(key, '')
                if value:
                    merge = True
                    tags.pop(x)
                    tags.pop(x)
                    match = "%s %s" % (t1[0], t2[0])
                    pos = value
                    tags.insert(x, (match, pos))
                    break
        matches = []
        for t in tags:
            if t[1] == "NNP" or t[1] == "NNI" or t[1] =="NNS":
                # if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NN":
                matches.append(t[0])
        return matches
    def extractSubstance(self):
        geo = self.geosubstance()
        tokens = self.tokenize_sentence(self.sentence)
        tags = self.normalize_tags(bigram_tagger.tag(tokens))
        merge = True
        while merge:
            merge = False
            for x in range(0, len(tags) - 1):
                t1 = tags[x]
                t2 = tags[x + 1]
                key = "%s+%s" % (t1[1], t2[1])
                value = cfg.get(key, '')
                if value:
                    merge = True
                    tags.pop(x)
                    tags.pop(x)
                    match = "%s %s" % (t1[0], t2[0])
                    pos = value
                    tags.insert(x, (match, pos))
                    break
        matches = []
        for t in tags:
            if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NNS":
                # if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NN":
                matches.append(t[0])
        result = []

        for m in range(len(matches)):
            for d in range(len(geo)):
                res = matches[m].lower().find(geo[d].lower())
                if (res != -1):
                    result.append(matches[m])

        correct = []
        for item in range(len(result)):
            if 'Ma' not in result[item]:
                correct.append(result[item])

        seen = set()
        Substance = []
        for item in correct:
            if item not in seen:
                seen.add(item)
                Substance.append(item)
        return Substance
    def extractTime(self):
        dict = self.dict()
        tokens = self.tokenize_sentence(self.sentence)
        tags = self.normalize_tags(bigram_tagger.tag(tokens))
        # print(tags)
        merge = True
        while merge:
            merge = False
            for x in range(0, len(tags) - 1):
                t1 = tags[x]
                t2 = tags[x + 1]
                key = "%s+%s" % (t1[1], t2[1])
                value = cfg.get(key, '')
                if value:
                    merge = True
                    tags.pop(x)
                    tags.pop(x)
                    match = "%s %s" % (t1[0], t2[0])
                    pos = value
                    tags.insert(x, (match, pos))
                    break
        matches = []
        for t in tags:
            if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NNS":
                matches.append(t[0])
        result = []
        for m in range(len(matches)):
            for d in range(len(dict)):
                res = matches[m].lower().find(dict[d].lower())
                if (res != -1):
                    result.append(matches[m])
        seen = set()
        Time = []
        for item in result:
            if item not in seen:
                seen.add(item)
                Time.append(item)
        return Time
def extractParagraph():
    countTxt = 0
    osfile = []
    dir = r"..\ToNeo"
    for root, dirs, files in os.walk(dir):
        for file in files:
            if file.split('.')[1] == 'txt':
                osfile.append(file)
                print(file)
    for index in range(len(osfile)):
        Geomiddle = []
        rew = open('../word/' + 'Aword.txt', "w", encoding='utf-8')
        re = open('../TimeSubstance/' + 'Substance_P' + osfile[index].split('.')[0] + '.txt', "w", encoding='utf-8')
        with open(osfile[index], "r", encoding='utf-8') as f:
            while True:
                data = f.readline()
                np_extractor = NPExtractor(data)
                result = np_extractor.extract()
                substance = np_extractor.extractSubstance()
                time = np_extractor.extractTime()
                # 进行每句切词分类
                substancefinall = []
                print("This result is about: %s" % ", ".join(result))
                for countsub in range(len(substance)):
                    count = len(substance[countsub].replace("0", " ").replace("1", " ").replace("2", " ").replace("3"," ").replace("4", " ")
                                .replace("5", " ").replace("6", " ").replace("7", " ").replace("8", " ").replace("9"," ").replace("\\"," ").replace("_"," ").replace("|"," ").replace("="," ").replace("—"," ")
                                .replace("^","").replace("©","").replace("<","").replace(">","").replace("'","").strip().split(" "))
                    countlen = len(substance[countsub].replace("0", " ").replace("1", " ").replace("2", " ").replace("3"," ").replace("4", " ")
                                .replace("5", " ").replace("6", " ").replace("7", " ").replace("8", " ").replace("9"," ").replace("\\"," ").replace("_"," ").replace("|"," ").replace("="," ").replace("—"," ")
                                   .replace("^","").replace("©","").replace("<","").replace(">","").replace("'","").strip())
                    if count < 4 and countlen > 4:
                        substancefinall.append(
                            substance[countsub].replace("-", " ").replace("~", " ").replace("/", " ").replace("0"," ").replace( "1", " ").replace("2", " ").replace("3", " ").replace("4", " ")
                            .replace("5", " ").replace("6", " ").replace("7", " ").replace("8", " ").replace("9"," ").replace("\\"," ").replace("_"," ").replace("|"," ").replace("="," ").replace("—"," ")
                                .replace("^","").replace("©","").replace("<","").replace(">","").replace("'","").strip().lower())
                        Geomiddle.append(
                            substance[countsub].replace("-", " ").replace("~", " ").replace("/", " ").replace("0"," ").replace("1", " ").replace("2", " ").replace("3", " ").replace("4", " ")
                            .replace("5", " ").replace("6", " ").replace("7", " ").replace("8", " ").replace("9"," ").replace("\\"," ").replace("_"," ").replace("|"," ").replace("="," ").replace("—"," ")
                                .replace("^","").replace("©","").replace("<","").replace(">","").replace("'","").strip().lower())
                # 分类写入txt
                if not data:
                    break
                re.write(data)
                re.write("Key words: %s" % ", ".join(result) + '\n')
                re.write("result: %s" % ", ".join(substancefinall) + '\r\n')
        re.close()
        for item in range(len(Geomiddle)):
            print(Geomiddle[item])
        seen = set()
        for item in Geomiddle:
            if item not in seen:
                seen.add(item)
                Geofinall.append(item)
        for windex in range(len(Geofinall)):
            rew.write(Geofinall[windex]+ '\n')
        rew.close()

def toNeo(py2neo=None):
    from py2neo import Graph, Node, Relationship, cypher, NodeMatcher, RelationshipMatcher
    graph = Graph("http://localhost:7474", auth=("neo4j", "admin"))
    matcher = NodeMatcher(graph)
    # 创建连接点
    # 创建节点之间的连接
    for index in range(len(Geofinall)):
        for hasindex in range(len(GeoDict)):
            if (Geofinall[index].lower().find(GeoDict[hasindex].lower()) != -1):
                # 创建节点
                if matcher.match(Geofinall[index].replace(" ", "_")).where(
                        "_.name=" + "\'" + Geofinall[index] + "\'").exists() == False:
                    loc_node = Node(Geofinall[index].replace(" ", "_"), label="outsource_"+Geofinall[index].replace(" ", "_"),
                                    name=Geofinall[index])
                    graph.create(loc_node)
                    print(Geofinall[index] + "节点不存在正在创建")
                    print(" ")
                    # 创建关系
                    print("======"+ GeoDict[hasindex] + "======")
                    has_node = matcher.match(GeoDict[hasindex].replace(" ", "_")).where(
                        "_.name=" + "\'" + GeoDict[hasindex] + "\'").first()
                    relationshipLevel = Relationship(has_node, 'subclass', loc_node)
                    print(Geofinall[index]+ " + " + GeoDict[hasindex]  + "边不存在正在创建")
                    graph.create(relationshipLevel)
                    print(Geofinall[index]+ " + " + GeoDict[hasindex]  + "关键创建成功")
                else:
                    loc_node = matcher.match(Geofinall[index].replace(" ", "_")).where(
                        "_.name=" + "\'" + Geofinall[index] + "\'").first()
                    print(Geofinall[index] + "节点已经存在")

if __name__ == '__main__':
    osfile = []
    Geofinall = []
    dir = r"D:\pythonProjectRelationshipToNeo4j\ToNeo"
    for root, dirs, files in os.walk(dir):
        for file in files:
            if file.split('.')[1] == 'docx':
                osfile.append(file)
    for index in range(len(osfile)):
        # 参数数组
        Dict = []
        GeoDict = []
        # 读取文本合并
        doc = Document(osfile[index])
        dict()
        geodict()
        getParagraphsText(doc,osfile[index])
        # 筛选
        filter()
        # 抽取
        extractParagraph()
        # 连接Neo4j
    toNeo()

